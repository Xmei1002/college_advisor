# app/services/volunteer/plan_service.py
from flask import current_app
from sqlalchemy.exc import SQLAlchemyError
from app.extensions import db
from app.models.student_volunteer_plan import StudentVolunteerPlan, VolunteerCollege, VolunteerSpecialty, VolunteerCategoryAnalysis
from app.services.college.recommendation_service import RecommendationService
from app.services.ai.moonshot import MoonshotAI
from app.services.ai.ollama import OllamaAPI
import json
from app.services.student.student_data_service import StudentDataService
from app.utils.helpers import convert_utc_to_beijing
class VolunteerPlanService:
    """志愿方案服务类，处理学生志愿方案相关业务逻辑"""
    
    @staticmethod
    def get_volunteer_plan(plan_id, include_details=True, category_id=None, group_id=None, volunteer_index=None):
        """
        获取志愿方案详情
        
        :param plan_id: 志愿方案ID
        :param include_details: 是否包含详细志愿信息
        :param category_id: 按类别ID过滤(1:冲, 2:稳, 3:保)
        :param group_id: 按志愿段ID过滤(1-12)
        :param volunteer_index: 按志愿序号过滤(1-48)
        :return: 志愿方案详情
        """
        plan = StudentVolunteerPlan.query.get_or_404(plan_id)
        result = plan.to_dict()
        
        # 获取志愿类别分析信息
        analyses_query = VolunteerCategoryAnalysis.query.filter_by(plan_id=plan_id)
        
        # 如果指定了类别ID，也按类别ID过滤分析
        if category_id is not None:
            analyses_query = analyses_query.filter_by(category_id=category_id)
            
        # 获取分析结果并添加到返回数据中
        category_analyses = analyses_query.all()
        result['category_analyses'] = [analysis.to_dict() for analysis in category_analyses]
        
        if include_details:
            # 构建基础查询
            colleges_query = VolunteerCollege.query.filter_by(plan_id=plan_id).order_by(VolunteerCollege.volunteer_index)

            # 应用过滤条件
            if category_id is not None:
                colleges_query = colleges_query.filter_by(category_id=category_id)
            if group_id is not None:
                colleges_query = colleges_query.filter_by(group_id=group_id)
            if volunteer_index is not None:
                colleges_query = colleges_query.filter_by(volunteer_index=volunteer_index)
                
            # 执行查询获取结果
            colleges = colleges_query.all()
            college_ids = [college.id for college in colleges]

            # 批量获取所有专业，避免N+1查询问题
            all_specialties = {}
            if college_ids:
                specialties = VolunteerSpecialty.query.filter(
                    VolunteerSpecialty.volunteer_college_id.in_(college_ids)
                ).all()
                
                # 按院校ID分组专业
                for specialty in specialties:
                    if specialty.volunteer_college_id not in all_specialties:
                        all_specialties[specialty.volunteer_college_id] = []
                    all_specialties[specialty.volunteer_college_id].append(specialty.to_dict())
            
            # 构建完整的志愿方案详情
            volunteer_colleges = []
            for college in colleges:
                college_dict = college.to_dict()
                college_dict['specialties'] = all_specialties.get(college.id, [])
                volunteer_colleges.append(college_dict)
            
            result['colleges'] = volunteer_colleges
        
        return result
        
    @staticmethod
    def update_volunteer_plan(plan_id, update_data):
        """
        修改学生志愿方案（创建新版本）
        
        :param plan_id: 当前志愿方案ID
        :param update_data: 更新数据，包含修改的院校和专业信息
        :return: 新的志愿方案
        """
        try:
            # 1. 获取当前方案和学生信息
            current_plan = StudentVolunteerPlan.query.get_or_404(plan_id)
            student_id = current_plan.student_id
            planner_id = current_plan.planner_id
            
            # 使用 with_for_update() 锁定查询，防止并发问题
            latest_plan = db.session.query(StudentVolunteerPlan).filter(
                StudentVolunteerPlan.student_id == student_id
            ).order_by(StudentVolunteerPlan.version.desc()).with_for_update().first()
            
            new_version = latest_plan.version + 1 if latest_plan else 1

            # 2. 将所有当前版本设为非当前版本
            db.session.query(StudentVolunteerPlan).filter(
                StudentVolunteerPlan.student_id == student_id,
                StudentVolunteerPlan.is_current == True
            ).update({"is_current": False}, synchronize_session=False)
            
            # 3. 创建新版本的方案
            new_plan = StudentVolunteerPlan(
                student_id=student_id,
                planner_id=planner_id,
                version=new_version,
                is_current=True,
                remarks=update_data.get('remarks', f"从版本{current_plan.version}修改"),
                generation_status=StudentVolunteerPlan.GENERATION_STATUS_SUCCESS,
                generation_progress=100,
                generation_message="手动修改志愿方案",
                user_data_hash=current_plan.user_data_hash,
                student_data_snapshot=current_plan.student_data_snapshot
            )
            db.session.add(new_plan)
            db.session.flush()  # 刷新以获取新ID
            new_plan_id = new_plan.id
            
            # 4. 识别前端修改了哪些批次(category_id和group_id组合)以及志愿组合键
            modified_batches = set()
            modified_combined_keys = set()
            for college in update_data.get('colleges', []):
                batch_key = (college.get('category_id'), college.get('group_id'))
                modified_batches.add(batch_key)
                combined_key = (college.get('group_id'), college.get('volunteer_index'))
                modified_combined_keys.add(combined_key)
            
            # 5. 处理未修改的批次 - 直接从当前版本复制
            if modified_batches:
                # 获取当前版本中所有院校数据
                current_colleges = VolunteerCollege.query.filter_by(plan_id=current_plan.id).all()
                
                # 找出未修改的批次的院校，且不在修改的组合键中
                colleges_to_copy = []
                for college in current_colleges:
                    combined_key = (college.group_id, college.volunteer_index)
                    if (college.category_id, college.group_id) not in modified_batches and combined_key not in modified_combined_keys:
                        # 复制院校数据
                        new_college = VolunteerCollege(
                            plan_id=new_plan_id,
                            category_id=college.category_id,
                            group_id=college.group_id,
                            volunteer_index=college.volunteer_index,
                            college_id=college.college_id,
                            college_name=college.college_name,
                            college_group_id=college.college_group_id,
                            score_diff=college.score_diff,
                            prediction_score=college.prediction_score,
                            recommend_type=college.recommend_type,
                            ai_analysis=college.ai_analysis,
                            area_name=college.area_name,
                            group_name=college.group_name,
                            min_tuition=college.min_tuition,
                            max_tuition=college.max_tuition,
                            min_score=college.min_score,
                            plan_number=college.plan_number,
                            school_type_text=college.school_type_text,
                            subject_requirements=college.subject_requirements,
                            tese_text=college.tese_text,
                            teshu_text=college.teshu_text,
                            uncode=college.uncode,
                            nature=college.nature
                        )
                        colleges_to_copy.append(new_college)
                
                # 批量添加未修改的院校
                if colleges_to_copy:
                    db.session.bulk_save_objects(colleges_to_copy)
                    db.session.flush()
                    
                    # 为复制的院校添加专业
                    original_college_ids = [c.id for c in current_colleges 
                                        if (c.category_id, c.group_id) not in modified_batches
                                        and (c.group_id, c.volunteer_index) not in modified_combined_keys]
                    
                    # 使用组合键作为映射
                    new_college_map = {(c.group_id, c.volunteer_index): c.id for c in VolunteerCollege.query.filter_by(plan_id=new_plan_id).all()}
                    
                    if original_college_ids:
                        # 获取原院校的所有专业
                        original_specialties = VolunteerSpecialty.query.filter(
                            VolunteerSpecialty.volunteer_college_id.in_(original_college_ids)
                        ).all()
                        
                        # 准备复制的专业数据
                        specialties_to_copy = []
                        for specialty in original_specialties:
                            # 找到对应的新院校ID
                            original_college = next((c for c in current_colleges if c.id == specialty.volunteer_college_id), None)
                            if original_college:
                                combined_key = (original_college.group_id, original_college.volunteer_index)
                                if combined_key in new_college_map:
                                    new_college_id = new_college_map[combined_key]
                                    
                                    # 复制专业数据
                                    new_specialty = VolunteerSpecialty(
                                        volunteer_college_id=new_college_id,
                                        specialty_id=specialty.specialty_id,
                                        specialty_code=specialty.specialty_code,
                                        specialty_name=specialty.specialty_name,
                                        specialty_index=specialty.specialty_index,
                                        prediction_score=specialty.prediction_score,
                                        plan_number=specialty.plan_number,
                                        tuition=specialty.tuition,
                                        remarks=specialty.remarks,
                                        ai_analysis=specialty.ai_analysis,
                                        fenshuxian_id=specialty.fenshuxian_id
                                    )
                                    specialties_to_copy.append(new_specialty)
                        
                        # 批量添加复制的专业
                        if specialties_to_copy:
                            db.session.bulk_save_objects(specialties_to_copy)
                            db.session.flush()
            
            # 6. 处理修改的批次 - 使用前端数据并补充完整信息
            if update_data.get('colleges'):
                # 导入必要的类
                from app.core.recommendation.repository import CollegeRepository
                from app.services.student.student_data_service import StudentDataService
                from app.models.zwh_xgk_fenzu_2025 import ZwhXgkFenzu2025
                from app.models.zwh_xgk_yuanxiao_2025 import ZwhXgkYuanxiao2025
                from app.models.zwh_xgk_fenshuxian_2025 import ZwhXgkFenshuxian2025
                from app.models.zwh_areas import ZwhAreas
                
                # 获取推荐信息用于查询完整数据
                student_data = StudentDataService.extract_college_recommendation_data(student_id)
                subject_type = int(student_data.get('subject_type') or 1)
                education_level = int(student_data.get('education_level') or 11)
                student_subjects = student_data.get('student_subjects', {})  # 获取学生选科情况
                
                # 收集需要查询的ID
                college_group_ids = []
                
                for college in update_data['colleges']:
                    college_group_ids.append(college.get('college_group_id'))
                
                # 查询院校组详细信息
                college_group_details = {}
                if college_group_ids:
                    # 查询院校详细信息和投档线信息
                    college_groups = db.session.query(
                        ZwhXgkFenzu2025.cgid,
                        ZwhXgkFenzu2025.minxuefei,
                        ZwhXgkFenzu2025.maxxuefei,
                        ZwhXgkFenzu2025.cgname,
                        ZwhXgkYuanxiao2025.tese,
                        ZwhXgkYuanxiao2025.leixing,
                        ZwhXgkYuanxiao2025.xingzhi,
                        ZwhXgkYuanxiao2025.teshu,
                        ZwhXgkYuanxiao2025.uncode,
                        ZwhAreas.aname.label('area_name'),
                        ZwhAreas.aid.label('area_id'),  # 添加地区ID用于获取完整路径
                        ZwhXgkFenshuxian2025.yuce.label('prediction_score'),
                        ZwhXgkFenshuxian2025.csbplannum.label('plan_number'),
                    ).join(
                        ZwhXgkYuanxiao2025,
                        ZwhXgkFenzu2025.newcid == ZwhXgkYuanxiao2025.cid
                    ).join(
                        ZwhAreas,
                        ZwhXgkYuanxiao2025.aid == ZwhAreas.aid
                    ).outerjoin(
                        ZwhXgkFenshuxian2025,
                        db.and_(
                            ZwhXgkFenshuxian2025.cgid == ZwhXgkFenzu2025.cgid,
                            ZwhXgkFenshuxian2025.spid == 32767,
                            ZwhXgkFenshuxian2025.suid == subject_type,
                            ZwhXgkFenshuxian2025.newbid == education_level
                        )
                    ).filter(
                        ZwhXgkFenzu2025.cgid.in_(college_group_ids)
                    ).all()
                    
                    # 处理院校类型、特色等文本
                    for group in college_groups:
                        tese_text = CollegeRepository.convert_code_to_text(group.tese, 'tese')
                        leixing_text = CollegeRepository.convert_code_to_text(group.leixing, 'leixing')
                        teshu_text = CollegeRepository.convert_code_to_text(group.teshu, 'teshu')
                        
                        # 获取完整地区路径
                        area_path = CollegeRepository.get_complete_area_path(group.area_id)
                        complete_area_name = ''.join([area['aname'] for area in area_path[1:]]) if len(area_path) > 1 else group.area_name
                        
                        # xingzhi为1表示公办，否则为民办
                        nature = '公办' if group.xingzhi == 1 else '民办'
                        
                        college_group_details[group.cgid] = {
                            'min_tuition': group.minxuefei,
                            'max_tuition': group.maxxuefei,
                            'group_name': group.cgname,
                            'area_name': complete_area_name,  # 使用完整地区名称
                            'tese_text': tese_text,
                            'leixing_text': leixing_text[0] if leixing_text else '',
                            'teshu_text': teshu_text,
                            'uncode': group.uncode,
                            'nature': nature,  # 使用正确的学校性质
                            'subject_requirements': student_subjects,  # 使用学生选科情况
                            'prediction_score': group.prediction_score,
                            'plan_number': group.plan_number
                        }
                
                # 处理专业详细信息 - 按(cgid, spid)组合查询
                specialty_details = {}
                
                # 对每个院校和专业组合进行查询
                for college in update_data['colleges']:
                    college_group_id = int(college.get('college_group_id', 0))
                    for specialty_data in college.get('specialties', []):
                        specialty_id = int(specialty_data.get('specialty_id', 0))
                        
                        # 查询此专业在此院校组的信息
                        specialty = db.session.query(
                            ZwhXgkFenshuxian2025.id.label('fenshuxian_id'),
                            ZwhXgkFenshuxian2025.tuitions,
                            ZwhXgkFenshuxian2025.yuce,
                            ZwhXgkFenshuxian2025.csbplannum
                        ).filter(
                            ZwhXgkFenshuxian2025.cgid == college_group_id,
                            ZwhXgkFenshuxian2025.spid == specialty_id,
                            ZwhXgkFenshuxian2025.suid == subject_type,
                            ZwhXgkFenshuxian2025.newbid == education_level
                        ).first()
                        
                        if specialty:
                            key = (college_group_id, specialty_id)
                            specialty_details[key] = {
                                'fenshuxian_id': specialty.fenshuxian_id,
                                'tuition': specialty.tuitions,
                                'prediction_score': specialty.yuce,
                                'plan_number': specialty.csbplannum
                            }
                
                # 处理修改的院校数据
                colleges_to_add = []
                
                for college_data in update_data['colleges']:
                    category_id = int(college_data.get('category_id', 0))
                    group_id = int(college_data.get('group_id', 0))
                    volunteer_index = int(college_data.get('volunteer_index', 0))
                    college_id = int(college_data.get('college_id', 0))
                    college_group_id = int(college_data.get('college_group_id', 0))
                    
                    # 获取补充信息
                    group_info = college_group_details.get(college_group_id, {})
                    
                    # 创建院校数据
                    college = VolunteerCollege(
                        plan_id=new_plan_id,
                        category_id=category_id,
                        group_id=group_id,
                        volunteer_index=volunteer_index,
                        college_id=college_id,
                        college_name=college_data.get('college_name', ''),
                        college_group_id=college_group_id,
                        score_diff=college_data.get('score_diff', 0),
                        prediction_score=group_info.get('prediction_score', 0),
                        recommend_type=VolunteerCollege.RECOMMEND_PLANNER,
                        min_tuition=group_info.get('min_tuition'),
                        max_tuition=group_info.get('max_tuition'),
                        min_score=group_info.get('prediction_score'),
                        plan_number=group_info.get('plan_number'),
                        area_name=group_info.get('area_name', ''),
                        group_name=group_info.get('group_name', ''),
                        school_type_text=group_info.get('leixing_text', ''),
                        subject_requirements=group_info.get('subject_requirements'),  # 添加选科要求
                        tese_text=group_info.get('tese_text'),
                        teshu_text=group_info.get('teshu_text'),
                        uncode=group_info.get('uncode', ''),
                        nature=group_info.get('nature', '')  # 使用正确的学校性质
                    )
                    
                    colleges_to_add.append(college)
                
                # 批量添加院校
                if colleges_to_add:
                    db.session.bulk_save_objects(colleges_to_add)
                    db.session.flush()
                    
                    # 查询新添加的院校志愿，获取它们的ID
                    college_id_map = {
                        (c.group_id, c.volunteer_index): c.id for c in VolunteerCollege.query.filter_by(plan_id=new_plan_id).all()
                    }
                    
                    # 记录日志，检查志愿序号映射
                    current_app.logger.info(f"院校ID映射: {college_id_map}")
                    
                    # 准备专业数据
                    specialties_to_add = []
                    processed_specialty_keys = set()  # 避免重复添加专业
                    
                    for college_data in update_data['colleges']:
                        volunteer_index = int(college_data.get('volunteer_index', 0))
                        group_id = int(college_data.get('group_id', 0))
                        college_group_id = int(college_data.get('college_group_id', 0))
                        
                        # 使用组合键查找院校ID
                        volunteer_key = (group_id, volunteer_index)
                        current_app.logger.info(f"处理院校志愿组合键: {volunteer_key}")
                        
                        if volunteer_key in college_id_map:
                            college_id = college_id_map[volunteer_key]
                            current_app.logger.info(f"找到对应院校ID: {college_id}")
                            
                            # 先删除该院校下所有现有专业
                            db.session.query(VolunteerSpecialty).filter(
                                VolunteerSpecialty.volunteer_college_id == college_id
                            ).delete(synchronize_session=False)
                            db.session.flush()
                            
                            for specialty_data in college_data.get('specialties', []):
                                specialty_id = int(specialty_data.get('specialty_id', 0))
                                specialty_index = int(specialty_data.get('specialty_index', 1))
                                
                                # 生成唯一键，避免处理重复专业
                                specialty_key = (college_id, specialty_index)
                                if specialty_key in processed_specialty_keys:
                                    continue
                                    
                                processed_specialty_keys.add(specialty_key)
                                
                                # 记录日志
                                current_app.logger.info(f"处理专业ID: {specialty_id}, 院校ID: {college_id}")
                                
                                # 使用(cgid, spid)组合键获取专业信息
                                cgid_spid_key = (college_group_id, specialty_id)
                                specialty_info = specialty_details.get(cgid_spid_key, {})
                                
                                specialty = VolunteerSpecialty(
                                    volunteer_college_id=college_id,
                                    specialty_id=specialty_id,
                                    specialty_code=specialty_data.get('specialty_code', ''),
                                    specialty_name=specialty_data.get('specialty_name', ''),
                                    specialty_index=specialty_index,
                                    prediction_score=specialty_info.get('prediction_score', 0),
                                    plan_number=specialty_info.get('plan_number', 0),
                                    tuition=specialty_info.get('tuition', 0),
                                    fenshuxian_id=specialty_info.get('fenshuxian_id', 0)
                                )
                                
                                specialties_to_add.append(specialty)
                        else:
                            current_app.logger.error(f"未找到志愿组合键 {volunteer_key} 的院校ID映射")
                    
                    # 批量添加专业
                    if specialties_to_add:
                        db.session.bulk_save_objects(specialties_to_add)
                        db.session.flush()
                
            # 7. 提交事务 - 只在所有操作完成后执行一次
            db.session.commit()
            
            # 8. 返回新版本的方案
            return VolunteerPlanService.get_volunteer_plan(new_plan_id)
            
        except Exception as e:
            # 任何异常都回滚事务
            db.session.rollback()
            current_app.logger.error(f"更新志愿方案失败: {str(e)}")
            raise    
    
    @staticmethod
    def create_empty_plan(student_id, planner_id, remarks, user_data_hash=None, 
                        generation_status='pending', generation_progress=0, 
                        generation_message=None,student_data_snapshot=None):
        """
        创建空的志愿方案
        
        :param student_id: 学生ID
        :param planner_id: 规划师ID
        :param remarks: 备注说明
        :param user_data_hash: 用户数据哈希
        :param generation_status: 生成状态
        :param generation_progress: 生成进度
        :param generation_message: 生成消息
        :return: 创建的志愿方案
        """
        try:
            # 获取最新版本号
            latest_plan = StudentVolunteerPlan.query.filter_by(
                student_id=student_id
            ).order_by(StudentVolunteerPlan.version.desc()).first()
            
            new_version = 1
            if latest_plan:
                new_version = latest_plan.version + 1
                
                # 将所有之前的方案设置为非当前版本
                db.session.query(StudentVolunteerPlan).filter(
                    StudentVolunteerPlan.student_id == student_id,
                    StudentVolunteerPlan.is_current == True
                ).update({"is_current": False}, synchronize_session=False)
            
            # 创建新的志愿方案
            plan = StudentVolunteerPlan(
                student_id=student_id,
                planner_id=planner_id,
                version=new_version,
                is_current=True,
                remarks=remarks,
                user_data_hash=user_data_hash,
                generation_status=generation_status,
                generation_progress=generation_progress,
                generation_message=generation_message or "开始生成志愿方案",
                student_data_snapshot=student_data_snapshot
            )
            
            db.session.add(plan)
            db.session.commit()
            
            return plan.to_dict()
        except SQLAlchemyError as e:
            db.session.rollback()
            current_app.logger.error(f"创建空志愿方案失败: {str(e)}")
            raise

    @staticmethod
    def add_volunteer_college(plan_id, college_data):
        """
        向志愿方案添加院校志愿
        
        :param plan_id: 志愿方案ID
        :param college_data: 院校志愿数据
        :return: 添加的院校志愿
        """
        try:
            # 检查方案是否存在
            # plan = StudentVolunteerPlan.query.get_or_404(plan_id)
            
            # 检查是否已存在相同序号的志愿
            existing = VolunteerCollege.query.filter_by(
                plan_id=plan_id, 
                volunteer_index=college_data.get('volunteer_index')
            ).first()
            
            if existing:
                # 可选择更新或抛出错误
                for key, value in college_data.items():
                    if key != 'specialties' and hasattr(existing, key):
                        setattr(existing, key, value)
                db.session.commit()
                return existing.to_dict()
            
            # 创建新的院校志愿
            college = VolunteerCollege(
                plan_id=plan_id,
                category_id=college_data.get('category_id'),
                group_id=college_data.get('group_id'),
                volunteer_index=college_data.get('volunteer_index'),
                college_id=college_data.get('college_id'),
                college_name=college_data.get('college_name'),
                college_group_id=college_data.get('college_group_id'),
                score_diff=college_data.get('score_diff'),
                prediction_score=college_data.get('prediction_score'),
                recommend_type=college_data.get('recommend_type', VolunteerCollege.RECOMMEND_AI),
                ai_analysis=college_data.get('ai_analysis', ''),
                area_name=college_data.get('area_name'),
                group_name=college_data.get('group_name'),
                min_tuition=college_data.get('min_tuition'),
                max_tuition=college_data.get('max_tuition'),
                min_score=college_data.get('min_score'),
                plan_number=college_data.get('plan_number'),
                school_type_text=college_data.get('school_type_text'),
                subject_requirements=college_data.get('subject_requirements'),
                tese_text=college_data.get('tese_text'),
                teshu_text=college_data.get('teshu_text'),
                uncode=college_data.get('uncode'),
                nature=college_data.get('nature'),
            )
            
            db.session.add(college)
            db.session.commit()
            
            return college.to_dict()
        except SQLAlchemyError as e:
            db.session.rollback()
            current_app.logger.error(f"添加院校志愿失败: {str(e)}")
            raise

    @staticmethod
    def add_volunteer_specialty(volunteer_college_id, specialty_data):
        """
        向院校志愿添加专业志愿
        
        :param volunteer_college_id: 院校志愿ID
        :param specialty_data: 专业志愿数据
        :return: 添加的专业志愿
        """
        try:
            # 检查院校志愿是否存在
            # college = VolunteerCollege.query.get_or_404(volunteer_college_id)
            
            # 检查是否已存在相同序号的专业
            existing = VolunteerSpecialty.query.filter_by(
                volunteer_college_id=volunteer_college_id, 
                specialty_index=specialty_data.get('specialty_index')
            ).first()
            
            if existing:
                # 可选择更新或抛出错误
                for key, value in specialty_data.items():
                    if hasattr(existing, key):
                        setattr(existing, key, value)
                db.session.commit()
                return existing.to_dict()
            
            # 创建新的专业志愿
            specialty = VolunteerSpecialty(
                volunteer_college_id=volunteer_college_id,
                specialty_id=specialty_data.get('specialty_id'),
                specialty_code=specialty_data.get('specialty_code'),
                specialty_name=specialty_data.get('specialty_name'),
                specialty_index=specialty_data.get('specialty_index'),
                prediction_score=specialty_data.get('prediction_score'),
                plan_number=specialty_data.get('plan_number'),
                tuition=specialty_data.get('tuition'),
                remarks=specialty_data.get('remarks', ''),
                ai_analysis=specialty_data.get('ai_analysis', ''),
                fenshuxian_id=specialty_data.get('fenshuxian_id')
            )
            
            db.session.add(specialty)
            db.session.commit()
            
            return specialty.to_dict()
        except SQLAlchemyError as e:
            db.session.rollback()
            current_app.logger.error(f"添加专业志愿失败: {str(e)}")
            raise

    @staticmethod
    def batch_add_volunteer_colleges(plan_id, colleges_data):
        """
        批量添加院校志愿
        
        :param plan_id: 志愿方案ID
        :param colleges_data: 院校志愿数据列表
        :return: 批量添加结果
        """
        try:
            current_app.logger.info(f"批量添加院校志愿, {len(colleges_data)} 个院校")
            
            # 获取现有院校志愿的序号
            existing_indexes = {c.volunteer_index for c in 
                            VolunteerCollege.query.filter_by(plan_id=plan_id).all()}
            
            colleges_to_add = []
            colleges_to_update = []
            
            # 分类需要添加和更新的院校志愿
            for college_data in colleges_data:
                # 确保所有数值字段都为有效的整数
                volunteer_index = int(college_data.get('volunteer_index', 0) or 0)
                category_id = int(college_data.get('category_id', 0) or 0)
                group_id = int(college_data.get('group_id', 0) or 0)
                college_id = int(college_data.get('college_id', 0) or 0)
                college_group_id = int(college_data.get('college_group_id', 0) or 0)
                score_diff = int(college_data.get('score_diff', 0) or 0)
                prediction_score = int(college_data.get('prediction_score', 0) or 0)
                
                if volunteer_index in existing_indexes:
                    # 获取现有院校志愿进行更新
                    college = VolunteerCollege.query.filter_by(
                        plan_id=plan_id, 
                        volunteer_index=volunteer_index
                    ).first()
                    
                    # 更新字段
                    college.category_id = category_id
                    college.group_id = group_id
                    college.college_id = college_id
                    college.college_name = college_data.get('college_name', '')
                    college.college_group_id = college_group_id
                    college.score_diff = score_diff
                    college.prediction_score = prediction_score
                    college.recommend_type = college_data.get('recommend_type', VolunteerCollege.RECOMMEND_AI)
                    college.ai_analysis = college_data.get('ai_analysis', ''),
                    college.area_name=college_data.get('area_name'),
                    college.group_name=college_data.get('group_name'),
                    college.min_tuition=college_data.get('min_tuition'),
                    college.max_tuition=college_data.get('max_tuition'),
                    college.min_score=college_data.get('min_score'),
                    college.plan_number=college_data.get('plan_number'),
                    college.school_type_text=college_data.get('school_type_text'),
                    college.subject_requirements=college_data.get('subject_requirements'),
                    college.tese_text=college_data.get('tese_text'),
                    college.teshu_text=college_data.get('teshu_text'),
                    college.uncode=college_data.get('uncode'),
                    college.nature=college_data.get('nature'),
                    
                    colleges_to_update.append(college)
                else:
                    # 创建新的院校志愿
                    college = VolunteerCollege(
                        plan_id=plan_id,
                        category_id=category_id,
                        group_id=group_id,
                        volunteer_index=volunteer_index,
                        college_id=college_id,
                        college_name=college_data.get('college_name', ''),
                        college_group_id=college_group_id,
                        score_diff=score_diff,
                        prediction_score=prediction_score,
                        recommend_type=college_data.get('recommend_type', VolunteerCollege.RECOMMEND_AI),
                        ai_analysis=college_data.get('ai_analysis'),
                        area_name=college_data.get('area_name'),
                        group_name=college_data.get('group_name'),
                        min_tuition=college_data.get('min_tuition'),
                        max_tuition=college_data.get('max_tuition'),
                        min_score=college_data.get('min_score'),
                        plan_number=college_data.get('plan_number'),
                        school_type_text=college_data.get('school_type_text'),
                        subject_requirements=college_data.get('subject_requirements'),
                        tese_text=college_data.get('tese_text'),
                        teshu_text=college_data.get('teshu_text'),
                        uncode=college_data.get('uncode'),
                        nature=college_data.get('nature'),
                    )
                    colleges_to_add.append(college)
            
            # 批量添加新院校志愿
            if colleges_to_add:
                db.session.bulk_save_objects(colleges_to_add)
            
            # 批量更新现有院校志愿
            if colleges_to_update:
                for college in colleges_to_update:
                    db.session.add(college)
            
            db.session.commit()
            
            return {
                'plan_id': plan_id,
                'added_count': len(colleges_to_add),
                'updated_count': len(colleges_to_update),
                'total_count': len(colleges_to_add) + len(colleges_to_update)
            }
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"批量添加院校志愿失败: {str(e)}")
            raise

    @staticmethod
    def batch_add_volunteer_specialties(volunteer_college_id, specialties_data):
        """
        批量添加专业志愿
        
        :param volunteer_college_id: 院校志愿ID
        :param specialties_data: 专业志愿数据列表
        :return: 批量添加结果
        """
        try:
            # 检查院校志愿是否存在
            # college = VolunteerCollege.query.get_or_404(volunteer_college_id)
            
            # 获取现有专业志愿的序号
            existing_indexes = {s.specialty_index for s in 
                              VolunteerSpecialty.query.filter_by(volunteer_college_id=volunteer_college_id).all()}
            
            specialties_to_add = []
            specialties_to_update = []
            
            # 分类需要添加和更新的专业志愿
            for specialty_data in specialties_data:
                index = specialty_data.get('specialty_index')
                if index in existing_indexes:
                    # 获取现有专业志愿进行更新
                    specialty = VolunteerSpecialty.query.filter_by(
                        volunteer_college_id=volunteer_college_id, 
                        specialty_index=index
                    ).first()
                    
                    for key, value in specialty_data.items():
                        if hasattr(specialty, key):
                            setattr(specialty, key, value)
                    
                    specialties_to_update.append(specialty)
                else:
                    # 创建新的专业志愿
                    specialty = VolunteerSpecialty(
                        volunteer_college_id=volunteer_college_id,
                        specialty_id=specialty_data.get('specialty_id'),
                        specialty_code=specialty_data.get('specialty_code'),
                        specialty_name=specialty_data.get('specialty_name'),
                        specialty_index=index,
                        prediction_score=specialty_data.get('prediction_score'),
                        plan_number=specialty_data.get('plan_number'),
                        tuition=specialty_data.get('tuition'),
                        remarks=specialty_data.get('remarks', ''),
                        ai_analysis=specialty_data.get('ai_analysis', ''),
                        fenshuxian_id=specialty_data.get('fenshuxian_id')
                    )
                    specialties_to_add.append(specialty)
            
            # 批量添加新专业志愿
            if specialties_to_add:
                db.session.bulk_save_objects(specialties_to_add)
            
            # 批量更新现有专业志愿
            if specialties_to_update:
                for specialty in specialties_to_update:
                    db.session.add(specialty)
            
            db.session.commit()
            
            return {
                'volunteer_college_id': volunteer_college_id,
                'added_count': len(specialties_to_add),
                'updated_count': len(specialties_to_update),
                'total_count': len(specialties_to_add) + len(specialties_to_update)
            }
            
        except SQLAlchemyError as e:
            db.session.rollback()
            current_app.logger.error(f"批量添加专业志愿失败: {str(e)}")
            raise

    @staticmethod
    def export_volunteer_plan_to_excel(plan_id):
        """
        将志愿方案导出为Excel文件，使用简单的斑马条纹样式，动态调整行高
        
        :param plan_id: 志愿方案ID
        :return: 生成的Excel文件路径
        """
        try:
            import os
            from datetime import datetime
            from flask import current_app
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            
            # 获取志愿方案详情
            plan_data = VolunteerPlanService.get_volunteer_plan(plan_id)
            if not plan_data or 'colleges' not in plan_data:
                return {
                    'success': False,
                    'error': "志愿方案数据不完整"
                }
            
            # 创建存储目录
            upload_folder = current_app.config.get('UPLOAD_FOLDER')
            if not upload_folder:
                # 如果配置中没有，使用默认路径
                upload_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'uploads')
                current_app.logger.warning(f"UPLOAD_FOLDER 配置缺失，使用默认路径: {upload_folder}")
                
            export_dir = os.path.join(upload_folder, 'exports')
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"volunteer_plan_{plan_id}_{timestamp}.xlsx"
            filepath = os.path.join(export_dir, filename)
            
            # 创建工作簿和工作表
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "志愿方案"
            
            # 定义样式
            header_fill = PatternFill(start_color="DDEBF7", end_color="DDEBF7", fill_type="solid")
            header_font = Font(bold=True)
            header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # 斑马条纹样式 - 简单的浅灰色
            light_gray_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            
            thin_border = Side(border_style="thin", color="000000")
            border = Border(left=thin_border, right=thin_border, top=thin_border, bottom=thin_border)
            
            # 设置左对齐但带有一定缩进的样式
            left_alignment_with_indent = Alignment(horizontal='left', vertical='center', wrap_text=True, indent=1)
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # 辅助函数：计算文本行数
            def count_lines(text):
                if not text:
                    return 1
                return text.count('\n') + 1
            
            # 辅助函数：根据内容计算行高
            def calculate_row_height(text, base_height=15, margin=10):
                lines = count_lines(text)
                return base_height * lines + margin
            
            # 添加表头
            headers = ["序号", "学校信息", "专业信息", "建议"]
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = header_alignment
                cell.border = border
            
            # 设置表头行高
            ws.row_dimensions[1].height = 25
            
            # 列宽设置
            ws.column_dimensions['A'].width = 10  # 序号列
            ws.column_dimensions['B'].width = 40  # 学校信息列
            ws.column_dimensions['C'].width = 60  # 专业信息列
            ws.column_dimensions['D'].width = 15  # 同意调剂列
            
            # 填充数据
            row_idx = 2
            for i, college in enumerate(plan_data['colleges']):
                # 使用简单的浅灰色/白色交替
                use_gray = (i % 2 == 0)
                
                # 根据志愿序号确定志愿名称
                volunteer_index = college.get('volunteer_index', 0)
                volunteer_name = f"第{volunteer_index}志愿"
                
                # 院校信息 - 为确保内容格式一致，使用每行两个空格的缩进
                school_info = f"名称: {college.get('college_name', '')}\n"
                school_info += f"代码: {college.get('college_id', '')}  性质: {college.get('nature', '公办')}\n"
                
                # 标签，如果有特色文本，用第一个特色作为标签
                tese_text = college.get('tese_text', [])
                tag = tese_text[0] if tese_text and isinstance(tese_text, list) and len(tese_text) > 0 else "硕博点"
                school_info += f"标签: {tag}\n"
                
                school_info += f"类型: {college.get('school_type_text', '')}\n"
                school_info += f"参考分数: {college.get('min_score', '')} \n"
                school_info += f"属地: {college.get('area_name', '')}\n"
                
                # 获取选科要求
                subject_reqs = college.get('subject_requirements', {})
                subjects = []
                if subject_reqs:
                    if subject_reqs.get('wu') == 1:
                        subjects.append("物理")
                    if subject_reqs.get('hua') == 1:
                        subjects.append("化学")
                    if subject_reqs.get('sheng') == 1:
                        subjects.append("生物")
                    if subject_reqs.get('shi') == 1:
                        subjects.append("历史")
                    if subject_reqs.get('di') == 1:
                        subjects.append("地理")
                    if subject_reqs.get('zheng') == 1:
                        subjects.append("政治")
                
                subject_text = "、".join(subjects) if subjects else "物理、化学"
                school_info += f"限报: {subject_text}"
                
                # 专业信息
                specialties = college.get('specialties', [])
                
                # 构建专业信息文本
                specialty_info = ""
                if specialties:
                    for sp_idx, specialty in enumerate(specialties):
                        specialty_name = specialty.get('specialty_name', '')
                        specialty_index = specialty.get('specialty_index', sp_idx + 1)
                        plan_number = specialty.get('plan_number', 0)
                        tuition = specialty.get('tuition', 0)

                        if sp_idx > 0:
                            specialty_info += "\n"
                        specialty_info += f"{specialty_index}: {specialty_name} [{plan_number} / {tuition}]"
                
                # 计算所需的行高 - 比较学校信息和专业信息，取较大者
                school_lines = count_lines(school_info)
                specialty_lines = count_lines(specialty_info) if specialty_info else 1
                max_lines = max(school_lines, specialty_lines)
                
                # 计算最终行高 - 每行16像素 + 10像素边距
                row_height = 16 * max_lines + 15
                
                # 设置行高
                ws.row_dimensions[row_idx].height = row_height
                
                # 序号单元格
                cell = ws.cell(row=row_idx, column=1)
                cell.value = volunteer_name
                cell.alignment = center_alignment
                cell.border = border
                if use_gray:
                    cell.fill = light_gray_fill
                
                # 学校信息单元格
                cell = ws.cell(row=row_idx, column=2)
                cell.value = school_info
                cell.alignment = left_alignment_with_indent
                cell.border = border
                if use_gray:
                    cell.fill = light_gray_fill
                
                # 专业信息单元格
                cell = ws.cell(row=row_idx, column=3)
                cell.value = specialty_info
                cell.alignment = left_alignment_with_indent
                cell.border = border
                if use_gray:
                    cell.fill = light_gray_fill
                
                # 同意调剂单元格
                # 根据类别确定建议
                category_id = college.get('category_id')
                if category_id:
                    recommendation = "冲" if category_id == 1 else ("稳" if category_id == 2 else "保")
                else:
                    recommendation = "冲"
                    
                cell = ws.cell(row=row_idx, column=4)
                cell.value = recommendation
                cell.alignment = center_alignment
                cell.border = border
                if use_gray:
                    cell.fill = light_gray_fill
                
                # 移动到下一行
                row_idx += 1
            
            # 保存Excel文件
            wb.save(filepath)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath,
                'url': f"/uploads/exports/{filename}"
            }
            
        except Exception as e:
            current_app.logger.error(f"导出志愿方案Excel失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    @staticmethod
    def export_volunteer_plan_analysis_to_excel(plan_id):
        """
        将志愿方案解析导出为Excel文件，包含整体解析、各类别解析和学校解析
        
        :param plan_id: 志愿方案ID
        :return: 生成的Excel文件路径信息
        """
        try:
            import os
            from datetime import datetime
            from flask import current_app
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from app.utils.helpers import convert_utc_to_beijing
            
            # 获取志愿方案详情（包含所有学校和解析信息）
            plan_data = VolunteerPlanService.get_volunteer_plan(plan_id)
            if not plan_data:
                return {
                    'success': False,
                    'error': "志愿方案数据不完整"
                }
            
            # 创建存储目录
            upload_folder = current_app.config.get('UPLOAD_FOLDER')
            if not upload_folder:
                # 如果配置中没有，使用默认路径
                upload_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'uploads')
                current_app.logger.warning(f"UPLOAD_FOLDER 配置缺失，使用默认路径: {upload_folder}")
                
            export_dir = os.path.join(upload_folder, 'exports')
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"volunteer_plan_analysis_{plan_id}_{timestamp}.xlsx"
            filepath = os.path.join(export_dir, filename)
            
            # 创建工作簿
            wb = openpyxl.Workbook()
            
            # 定义样式
            title_font = Font(name='微软雅黑', size=16, bold=True, color='4F81BD')
            subtitle_font = Font(name='微软雅黑', size=12, bold=True, color='4F81BD')
            header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
            content_font = Font(name='微软雅黑', size=10)
            bold_font = Font(name='微软雅黑', size=10, bold=True)
            italic_font = Font(name='微软雅黑', size=10, italic=True)
            
            title_alignment = Alignment(horizontal='center', vertical='center')
            subtitle_alignment = Alignment(horizontal='left', vertical='center')
            header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            content_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            odd_row_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            
            thin_border = Side(border_style="thin", color="000000")
            border = Border(left=thin_border, right=thin_border, top=thin_border, bottom=thin_border)
            
            # 颜色定义
            colors = {
                '冲': 'FF9999',  # 浅红色
                '稳': 'FFCC99',  # 浅橙色
                '保': '99CC99'   # 浅绿色
            }

            # 辅助函数：处理Markdown格式文本
            def format_markdown_text(cell, text):
                """
                简单处理Markdown格式文本，转换为富文本
                支持基本的加粗、斜体、标题、列表等
                """
                if not text:
                    return
                
                # 最简单的处理方式：去除Markdown标记但保留文本内容
                # 实际应用中可以进一步区分样式
                cell.value = text.replace('**', '').replace('*', '').replace('#', '').replace('`', '')
                
                # 更复杂的处理可以在这里添加，例如识别不同标记并应用不同样式
                # 这里只是示例，实际实现可能需要更复杂的解析
            
            # 1. 创建总体概览工作表
            ws_overview = wb.active
            ws_overview.title = "总体方案解析"
            
            # 设置列宽
            ws_overview.column_dimensions['A'].width = 15
            ws_overview.column_dimensions['B'].width = 70
            
            # 标题
            ws_overview.merge_cells('A1:B1')
            cell = ws_overview['A1']
            cell.value = f"志愿方案分析报告"
            cell.font = title_font
            cell.alignment = title_alignment
            ws_overview.row_dimensions[1].height = 30
            
            # 基本信息
            row = 3
            ws_overview.cell(row=row, column=1).value = "方案名称"
            ws_overview.cell(row=row, column=1).font = subtitle_font
            ws_overview.cell(row=row, column=1).alignment = subtitle_alignment
            
            ws_overview.cell(row=row, column=2).value = plan_data.get('name', '未命名方案')
            ws_overview.cell(row=row, column=2).font = content_font
            ws_overview.cell(row=row, column=2).alignment = content_alignment
            
            row += 1
            ws_overview.cell(row=row, column=1).value = "创建时间"
            ws_overview.cell(row=row, column=1).font = subtitle_font
            ws_overview.cell(row=row, column=1).alignment = subtitle_alignment
            
            # 获取并转换UTC时间为北京时间
            created_at = plan_data.get('created_at', '')
            created_at = convert_utc_to_beijing(created_at)
            
            ws_overview.cell(row=row, column=2).value = created_at
            ws_overview.cell(row=row, column=2).font = content_font
            ws_overview.cell(row=row, column=2).alignment = content_alignment
            
            row += 2
            ws_overview.cell(row=row, column=1).value = "整体解析"
            ws_overview.cell(row=row, column=1).font = subtitle_font
            ws_overview.cell(row=row, column=1).alignment = subtitle_alignment
            
            row += 1
            # 查找category_id为0的整体解析
            overall_analysis = '暂无整体解析'
            category_analyses = plan_data.get('category_analyses', [])
            for analysis in category_analyses:
                if analysis.get('category_id') == 0:
                    overall_analysis = analysis.get('analysis_content', '暂无整体解析')
                    break
                    
            ws_overview.merge_cells(f'A{row}:B{row+5}')
            cell = ws_overview[f'A{row}']
            # 处理可能的Markdown格式
            format_markdown_text(cell, overall_analysis)
            cell.font = content_font
            cell.alignment = content_alignment
            
            # 2. 创建分类解析工作表
            ws_category = wb.create_sheet(title="分类解析")
            
            # 设置列宽
            ws_category.column_dimensions['A'].width = 15
            ws_category.column_dimensions['B'].width = 60
            
            # 添加标题
            ws_category.merge_cells('A1:B1')
            cell = ws_category['A1']
            cell.value = "志愿方案分类解析"
            cell.font = title_font
            cell.alignment = title_alignment
            ws_category.row_dimensions[1].height = 30
            
            # 表头
            row = 3
            headers = ["类别", "分析"]
            for col_idx, header in enumerate(headers, 1):
                cell = ws_category.cell(row=row, column=col_idx)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            # 填充分类数据
            row += 1
            categories = {'1': '冲', '2': '稳', '3': '保'}
            category_analyses = plan_data.get('category_analyses', [])
            
            for category_analysis in category_analyses:
                category_id = category_analysis.get('category_id')
                # 跳过整体解析（category_id=0）
                if category_id == 0:
                    continue
                    
                category_name = categories.get(str(category_id), '未知')
                analysis = category_analysis.get('analysis_content', '暂无分析')
                
                # 设置颜色填充
                fill = PatternFill(start_color=colors.get(category_name, "FFFFFF"), 
                                   end_color=colors.get(category_name, "FFFFFF"), 
                                   fill_type="solid")
                
                # 添加数据行
                cell = ws_category.cell(row=row, column=1)
                cell.value = category_name
                cell.font = content_font
                cell.alignment = center_alignment
                cell.border = border
                cell.fill = fill
                
                cell = ws_category.cell(row=row, column=2)
                # 处理可能的Markdown格式
                format_markdown_text(cell, analysis)
                cell.font = content_font
                cell.alignment = content_alignment
                cell.border = border
                cell.fill = fill
                
                row += 1
            
            # 3. 创建学校解析工作表
            ws_school = wb.create_sheet(title="院校解析")
            
            # 设置列宽
            ws_school.column_dimensions['A'].width = 10
            ws_school.column_dimensions['B'].width = 30
            ws_school.column_dimensions['C'].width = 15
            ws_school.column_dimensions['D'].width = 60
            
            # 添加标题
            ws_school.merge_cells('A1:D1')
            cell = ws_school['A1']
            cell.value = "院校解析详情"
            cell.font = title_font
            cell.alignment = title_alignment
            ws_school.row_dimensions[1].height = 30
            
            # 表头
            row = 3
            headers = ["志愿序号", "院校名称", "类别", "AI解析"]
            for col_idx, header in enumerate(headers, 1):
                cell = ws_school.cell(row=row, column=col_idx)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            # 填充学校数据
            row += 1
            colleges = plan_data.get('colleges', [])
            
            for i, college in enumerate(colleges):
                volunteer_index = college.get('volunteer_index', 0)
                college_name = college.get('college_name', '')
                category_id = college.get('category_id', '')
                category_name = categories.get(str(category_id), '未知')
                ai_analysis = college.get('ai_analysis', '')
                
                # 跳过没有AI解析的学校
                if not ai_analysis:
                    continue
                
                # 设置行背景色 - 使用浅色斑马条纹
                is_odd_row = (i % 2 == 0)
                fill = odd_row_fill if is_odd_row else None
                
                # 添加数据行
                cell = ws_school.cell(row=row, column=1)
                cell.value = f"第{volunteer_index}志愿"
                cell.font = content_font
                cell.alignment = center_alignment
                cell.border = border
                if fill:
                    cell.fill = fill
                
                cell = ws_school.cell(row=row, column=2)
                cell.value = college_name
                cell.font = content_font
                cell.alignment = content_alignment
                cell.border = border
                if fill:
                    cell.fill = fill
                
                # 设置类别单元格的填充颜色
                cell = ws_school.cell(row=row, column=3)
                cell.value = category_name
                cell.font = content_font
                cell.alignment = center_alignment
                cell.border = border
                cat_fill = PatternFill(start_color=colors.get(category_name, "FFFFFF"), 
                                      end_color=colors.get(category_name, "FFFFFF"), 
                                      fill_type="solid")
                cell.fill = cat_fill
                
                cell = ws_school.cell(row=row, column=4)
                # 处理可能的Markdown格式
                format_markdown_text(cell, ai_analysis)
                cell.font = content_font
                cell.alignment = content_alignment
                cell.border = border
                if fill:
                    cell.fill = fill
                
                row += 1
            
            # 保存Excel文件
            wb.save(filepath)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath,
                'url': f"/uploads/exports/{filename}"
            }
            
        except Exception as e:
            current_app.logger.error(f"导出志愿方案解析Excel失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    @staticmethod
    def export_volunteer_plan_analysis_to_word(plan_id):
        """
        将志愿方案解析导出为Word文件，包含整体解析、各类别解析和学校解析
        支持Markdown格式的解析内容
        
        :param plan_id: 志愿方案ID
        :return: 生成的Word文件路径信息
        """
        try:
            import os
            from datetime import datetime
            from flask import current_app
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from app.utils.helpers import convert_utc_to_beijing
            
            # 获取志愿方案详情
            plan_data = VolunteerPlanService.get_volunteer_plan(plan_id)
            if not plan_data:
                return {
                    'success': False,
                    'error': "志愿方案数据不完整"
                }
            
            # 创建存储目录
            upload_folder = current_app.config.get('UPLOAD_FOLDER')
            if not upload_folder:
                upload_folder = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'uploads')
                current_app.logger.warning(f"UPLOAD_FOLDER 配置缺失，使用默认路径: {upload_folder}")
                
            export_dir = os.path.join(upload_folder, 'exports')
            os.makedirs(export_dir, exist_ok=True)
            
            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"volunteer_plan_analysis_{plan_id}_{timestamp}.docx"
            filepath = os.path.join(export_dir, filename)
            
            # 创建文档
            doc = Document()
            
            # 设置文档页面和字体
            sections = doc.sections
            for section in sections:
                section.page_width = Cm(21)  # A4宽度
                section.page_height = Cm(29.7)  # A4高度
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
            
            # 辅助函数：处理Markdown格式文本添加到Word
            def format_markdown_to_word(doc, text):
                """
                直接接收文档对象，处理Markdown文本并添加到文档，
                确保标题出现在相应内容之前
                """
                if not text:
                    return
                
                # 解析文本，将其分成标题和段落
                sections = []  # 存储解析后的部分，每部分包含标题和内容
                current_heading = None
                current_content = []
                
                lines = text.split('\n')
                for line in lines:
                    if line.startswith('#'):  # 是标题行
                        # 如果已有内容，保存当前部分
                        if current_heading or current_content:
                            sections.append((current_heading, current_content))
                            current_content = []
                        
                        # 保存新标题
                        current_heading = line
                    else:
                        # 添加到当前内容
                        current_content.append(line)
                
                # 添加最后一部分
                if current_heading or current_content:
                    sections.append((current_heading, current_content))
                
                # 按顺序添加到文档
                for heading, content in sections:
                    # 先添加标题
                    if heading:
                        if heading.startswith('# '):
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run(heading[2:])
                            run.font.size = Pt(16)
                            run.font.bold = True
                        elif heading.startswith('## '):
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run(heading[3:])
                            run.font.size = Pt(14)
                            run.font.bold = True
                        elif heading.startswith('### '):
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run(heading[4:])
                            run.font.size = Pt(13)
                            run.font.bold = True
                    
                    # 再添加内容
                    if content:
                        p = doc.add_paragraph()
                        # 处理内容中的格式
                        for line in content:
                            # 处理空行
                            if not line.strip():
                                p.add_run('\n')
                                continue
                            
                            # 处理粗体、斜体等
                            segments = []
                            current_pos = 0
                            in_bold = False
                            in_italic = False
                            
                            for i in range(len(line)):
                                if i < current_pos:
                                    continue
                                
                                if line[i:i+2] == '**' and not in_italic:
                                    segments.append((line[current_pos:i], in_bold, in_italic))
                                    current_pos = i + 2
                                    in_bold = not in_bold
                                elif line[i:i+1] == '*' and not in_bold:
                                    segments.append((line[current_pos:i], in_bold, in_italic))
                                    current_pos = i + 1
                                    in_italic = not in_italic
                            
                            # 添加最后一段文本
                            if current_pos < len(line):
                                segments.append((line[current_pos:], in_bold, in_italic))
                            
                            # 创建格式化的文本
                            for text, bold, italic in segments:
                                if not text:
                                    continue
                                run = p.add_run(text)
                                run.font.size = Pt(11)
                                run.font.bold = bold
                                run.italic = italic
                            
                            p.add_run('\n')
            
            # 1. 添加标题
            title = doc.add_heading('志愿方案分析报告', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 2. 添加基本信息
            doc.add_paragraph().add_run('').add_break()  # 添加空行
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run('创建时间：')
            run.font.bold = True
            run.font.size = Pt(12)
            
            # 获取并转换UTC时间为北京时间
            created_at = plan_data.get('created_at', '')
            created_at = convert_utc_to_beijing(created_at)
            
            p.add_run(created_at).font.size = Pt(11)
            
            # 3. 添加整体解析
            doc.add_paragraph().add_run('').add_break()  # 添加空行
            p = doc.add_paragraph()
            run = p.add_run('整体解析')
            run.font.bold = True
            run.font.size = Pt(14)
            
            # 查找category_id为0的整体解析
            overall_analysis = '暂无整体解析'
            category_analyses = plan_data.get('category_analyses', [])
            for analysis in category_analyses:
                if analysis.get('category_id') == 0:
                    overall_analysis = analysis.get('analysis_content', '暂无整体解析')
                    break
            
            # 使用修改后的函数处理Markdown
            format_markdown_to_word(doc, overall_analysis)
            
            doc.add_paragraph().add_run('').add_break()  # 添加空行
            
            # 4. 添加分类解析
            doc.add_heading('志愿方案分类解析', level=1)
            
            categories = {'1': '冲', '2': '稳', '3': '保'}
            category_colors = {
                '冲': RGBColor(255, 153, 153),  # 浅红色
                '稳': RGBColor(255, 204, 153),  # 浅橙色
                '保': RGBColor(153, 204, 153)   # 浅绿色
            }
            
            for category_analysis in category_analyses:
                category_id = category_analysis.get('category_id')
                # 跳过整体解析（category_id=0）
                if category_id == 0:
                    continue
                    
                category_name = categories.get(str(category_id), '未知')
                analysis = category_analysis.get('analysis_content', '暂无分析')
                
                p = doc.add_paragraph()
                run = p.add_run(f"{category_name}类院校解析")
                run.font.bold = True
                run.font.size = Pt(13)
                # 设置分类颜色
                if category_name in category_colors:
                    run.font.color.rgb = category_colors[category_name]
                
                # 使用修改后的函数处理Markdown
                format_markdown_to_word(doc, analysis)
                
                doc.add_paragraph().add_run('').add_break()  # 添加空行
            
            # 5. 添加院校解析
            doc.add_heading('院校解析详情', level=1)
            
            colleges = plan_data.get('colleges', [])
            has_ai_analysis = False
            
            for college in colleges:
                volunteer_index = college.get('volunteer_index', 0)
                college_name = college.get('college_name', '')
                category_id = college.get('category_id', '')
                category_name = categories.get(str(category_id), '未知')
                ai_analysis = college.get('ai_analysis', '')
                
                # 跳过没有AI解析的学校
                if not ai_analysis:
                    continue
                
                has_ai_analysis = True
                
                p = doc.add_paragraph()
                run = p.add_run(f"第{volunteer_index}志愿 - {college_name}")
                run.font.bold = True
                run.font.size = Pt(12)
                
                p = doc.add_paragraph()
                run = p.add_run(f"类别: {category_name}")
                run.font.size = Pt(11)
                if category_name in category_colors:
                    run.font.color.rgb = category_colors[category_name]
                
                # 使用修改后的函数处理Markdown
                format_markdown_to_word(doc, ai_analysis)
                
                doc.add_paragraph().add_run('').add_break()  # 添加空行
            
            if not has_ai_analysis:
                p = doc.add_paragraph()
                p.add_run('暂无院校AI解析数据').font.italic = True
            
            # 保存文档
            doc.save(filepath)
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath,
                'url': f"/uploads/exports/{filename}"
            }
            
        except Exception as e:
            current_app.logger.error(f"导出志愿方案解析Word失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
        
    # @staticmethod
    # def export_volunteer_plan_analysis_to_pdf(plan_id):
    #     """
    #     将志愿方案解析导出为PDF文件，使用python-docx读取Word内容，ReportLab生成PDF
    #     支持中文字符
        
    #     :param plan_id: 志愿方案ID
    #     :return: 生成的PDF文件路径信息
    #     """
    #     try:
    #         import os
    #         import platform
    #         from flask import current_app
    #         import docx
    #         from reportlab.lib.pagesizes import A4
    #         from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    #         from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    #         from reportlab.pdfbase import pdfmetrics
    #         from reportlab.pdfbase.ttfonts import TTFont
            
    #         # 先生成Word文档
    #         word_result = VolunteerPlanService.export_volunteer_plan_analysis_to_word(plan_id)
    #         if not word_result['success']:
    #             return word_result
            
    #         # 获取Word文件路径
    #         word_filepath = word_result['filepath']
            
    #         # 生成PDF文件名和路径
    #         filename = os.path.splitext(os.path.basename(word_filepath))[0] + '.pdf'
    #         export_dir = os.path.dirname(word_filepath)
    #         filepath = os.path.join(export_dir, filename)
            
    #         # 注册中文字体 - 方法一：使用嵌入式字体
    #         # 获取项目根目录
    #         project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    #         font_dir = os.path.join(project_root, 'app', 'static', 'fonts')
    #         os.makedirs(font_dir, exist_ok=True)
            
    #         # 使用中文字体，优先尝试嵌入式字体
    #         font_path = os.path.join(font_dir, 'SimSun.ttf')  # 宋体
            
    #         # 如果嵌入式字体不存在，尝试使用系统字体
    #         if not os.path.exists(font_path):
    #             system = platform.system().lower()
    #             if system == 'windows':
    #                 # Windows 系统字体路径
    #                 system_font_path = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'simsun.ttc')
    #                 if os.path.exists(system_font_path):
    #                     font_path = system_font_path
    #                 else:
    #                     # 备选字体
    #                     system_font_path = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'msyh.ttc')
    #                     if os.path.exists(system_font_path):
    #                         font_path = system_font_path
    #             elif system == 'linux':
    #                 # Linux 系统尝试常见的中文字体
    #                 linux_fonts = [
    #                     '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # 文泉驿正黑
    #                     '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',  # Droid Sans
    #                     '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc'  # Noto Sans CJK
    #                 ]
    #                 for f in linux_fonts:
    #                     if os.path.exists(f):
    #                         font_path = f
    #                         break
    #             elif system == 'darwin':  # macOS
    #                 mac_fonts = [
    #                     '/Library/Fonts/PingFang.ttc',  # PingFang
    #                     '/Library/Fonts/STHeiti Light.ttc',  # Heiti
    #                     '/System/Library/Fonts/PingFang.ttc'  # 系统 PingFang
    #                 ]
    #                 for f in mac_fonts:
    #                     if os.path.exists(f):
    #                         font_path = f
    #                         break
            
    #         # 如果找不到任何中文字体，使用备选方案：从在线下载或使用包内置字体
    #         if not os.path.exists(font_path):
    #             current_app.logger.warning("未找到中文字体，将尝试使用默认字体")
    #             # 这里可以添加从包内复制内置字体的逻辑，或者从在线下载字体
    #             # 简单起见，我们这里使用 Helvetica（不支持中文，但至少不会报错）
    #             font_name = "Helvetica"
    #         else:
    #             # 注册找到的字体
    #             font_name = "SimSun"  # 注册的字体名称
    #             pdfmetrics.registerFont(TTFont(font_name, font_path))
    #             current_app.logger.info(f"已注册中文字体: {font_path}")
            
    #         # 读取Word文档内容
    #         doc = docx.Document(word_filepath)
            
    #         # 创建PDF文档
    #         pdf_doc = SimpleDocTemplate(
    #             filepath,
    #             pagesize=A4,
    #             rightMargin=72,
    #             leftMargin=72,
    #             topMargin=72,
    #             bottomMargin=72
    #         )
            
    #         # 创建自定义样式
    #         styles = getSampleStyleSheet()
            
    #         # 修改默认样式，使用中文字体
    #         title_style = ParagraphStyle(
    #             'ChineseTitle',
    #             parent=styles['Title'],
    #             fontName=font_name,
    #             fontSize=20,
    #             leading=24
    #         )
            
    #         heading1_style = ParagraphStyle(
    #             'ChineseHeading1',
    #             parent=styles['Heading1'],
    #             fontName=font_name,
    #             fontSize=16,
    #             leading=20
    #         )
            
    #         normal_style = ParagraphStyle(
    #             'ChineseNormal',
    #             parent=styles['Normal'],
    #             fontName=font_name,
    #             fontSize=12,
    #             leading=16
    #         )
            
    #         # 转换内容
    #         elements = []
    #         for para in doc.paragraphs:
    #             text = para.text.strip()
    #             if not text:
    #                 continue
                    
    #             # 处理各种样式
    #             if para.style.name == 'Title' or para.style.name.startswith('Title'):
    #                 elements.append(Paragraph(text, title_style))
    #             elif para.style.name.startswith('Heading') or para.style.name == 'Heading 1':
    #                 elements.append(Paragraph(text, heading1_style))
    #             else:
    #                 # 处理文本中的特殊XML字符
    #                 text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    #                 elements.append(Paragraph(text, normal_style))
                
    #             elements.append(Spacer(1, 12))
            
    #         # 生成PDF
    #         pdf_doc.build(elements)
            
    #         # 检查是否成功生成
    #         if os.path.exists(filepath):
    #             return {
    #                 'success': True,
    #                 'filename': filename,
    #                 'filepath': filepath,
    #                 'url': f"/uploads/exports/{filename}"
    #             }
    #         else:
    #             return {
    #                 'success': False,
    #                 'error': "PDF生成失败，文件不存在",
    #                 'word_file': word_result
    #             }
            
    #     except Exception as e:
    #         current_app.logger.error(f"导出志愿方案解析PDF失败: {str(e)}")
    #         return {
    #             'success': False,
    #             'error': str(e),
    #             'word_file': word_result if 'word_result' in locals() else None
    #         }


def ai_select_college_ids(filtered_colleges, user_info, recommendation_data, is_first=False):
    """
    AI选择院校ID，基于学生成绩类型选择不同的推荐策略
    
    :param filtered_colleges: 筛选出的院校列表(包含简化数据)
    :param user_info: 用户信息文本
    :param recommendation_data: 学生推荐数据，包含各类成绩信息
    :return: 选择的院校ID列表
    """
    try:
        # 检查学生是否有高考成绩
        has_gaokao_score = recommendation_data.get('student_score', 0) > 0
        
        # 检查学生是否有模考成绩
        latest_mock_score = recommendation_data.get('mock_exam_score', 0) > 0
        if is_first:
            return fallback_recommendation(filtered_colleges)

        # 基于成绩类型选择推荐策略
        if has_gaokao_score:
            # 有高考成绩且不是第一次生成方案，使用AI推荐
            current_app.logger.info("学生有高考成绩，使用AI推荐")
            return ai_recommend_with_score(filtered_colleges, user_info)
        elif latest_mock_score:

            return fallback_recommendation(filtered_colleges)
        else:
            # 没有任何成绩，使用备选方案
            current_app.logger.info("学生没有可用成绩，使用备选推荐方案")
            return fallback_recommendation(filtered_colleges)
        
    except Exception as e:
        # 任何错误时使用备选方案
        current_app.logger.error(f"AI选择院校ID过程中发生错误: {str(e)}")
        return fallback_recommendation(filtered_colleges)

def ai_recommend_with_score(filtered_colleges, user_info):
    """使用AI基于高考成绩推荐院校"""
    # 这里是原来的AI推荐逻辑
    simplified_colleges = []
    for college in filtered_colleges:
        simplified_colleges.append({
            'cgid': college['cgid'],  # 院校专业组ID
            'name': college['cname'],  # 院校名称
            'city': college['area_name'],  # 城市
            'tese': college['tese_text'],  # 院校特色
            'specialties': [
                {
                    "spname": specialty["spname"],
                    "tuition": specialty["tuition"],
                    "spid": specialty["spid"]
                }
                for specialty in college['specialties']
            ]
        })
    
    simplified_colleges_json = json.dumps(simplified_colleges, ensure_ascii=False)
    # 调用AI服务并解析结果
    ai_res_json = MoonshotAI.filter_colleges(user_info, simplified_colleges_json)
    # ai_res_json = OllamaAPI.filter_colleges(user_info, simplified_colleges_json)
    ai_res_dict = json.loads(ai_res_json)
    
    # 验证AI返回结果是否符合要求（最多4个学校）
    if len(ai_res_dict) > 4:
        # 只保留前4个学校
        ai_res_dict = {k: ai_res_dict[k] for k in list(ai_res_dict.keys())[:4]}
    
    # 验证每个学校的专业是否符合要求（最多6个专业）
    for cgid, spids in ai_res_dict.items():
        if len(spids) > 6:
            ai_res_dict[cgid] = spids[:6]
    
    return ai_res_dict

def fallback_recommendation(filtered_colleges):
    """备选推荐方案"""
    # 准备备选方案（前4个学校，每个学校前6个专业）
    fallback_result = {}
    for college in filtered_colleges[:4]:
        cgid = str(college['cgid'])
        specialties = college['specialties'][:6]
        if specialties:
            fallback_result[cgid] = [str(s['spid']) for s in specialties]
    
    return fallback_result
    
def process_batch(student_id, planner_id, category_id, group_id, plan_id=None, is_first=False):
    """
    处理一个批次的志愿
    
    :param student_id: 学生ID
    :param planner_id: 规划师ID
    :param category_id: 类别ID(1:冲, 2:稳, 3:保)
    :param group_id: 组内ID(1-4)
    :param plan_id: 志愿方案ID，如果已有
    :return: 更新后的志愿方案ID和批次处理状态
    """
    
    try:
        # 计算实际的group_id（1-12）
        actual_group_id = (category_id - 1) * 4 + group_id

        current_app.logger.info(f"=====类别ID={category_id}, 分组ID={actual_group_id}, 方案ID={plan_id}=====")
        # 使用StudentDataService获取学生数据
        recommendation_data = StudentDataService.extract_college_recommendation_data(student_id)
        # 获取学生的文本信息
        user_info = StudentDataService.generate_student_profile_text(student_id)

        student_score = int(recommendation_data['student_score'] or 0)
        subject_type = int(recommendation_data['subject_type'] or 1)
        education_level = int(recommendation_data['education_level'] or 11)
        student_subjects = recommendation_data['student_subjects']
        area_ids = recommendation_data.get('area_ids', [])
        specialty_types = recommendation_data.get('specialty_types', [])
        tuition_ranges = recommendation_data.get('tuition_ranges', [])
        mock_exam_score = int(recommendation_data['mock_exam_score'] or 0)

        # 确保area_ids和specialty_types是可迭代的且包含有效整数
        area_ids = [int(aid) for aid in area_ids if aid and str(aid).isdigit()]
        specialty_types = [int(st) for st in specialty_types if st and str(st).isdigit()]
        
        # 1. 获取筛选结果
        filtered_colleges, _ = RecommendationService.get_colleges_by_category_and_group(
            student_score=student_score > 0 and student_score or mock_exam_score,
            subject_type=subject_type,
            education_level=education_level,
            category_id=category_id,
            group_id=actual_group_id,
            student_subjects=student_subjects,
            area_ids=area_ids,
            specialty_types=specialty_types,
            tuition_ranges=tuition_ranges,
            page=1,
            per_page=100  # 获取足够多的结果供AI选择
        )
        
        current_app.logger.info(f"筛选到的院校数量: {len(filtered_colleges)}")
        
        # 如果没有筛选到院校，直接返回
        if not filtered_colleges:
            return plan_id, False
        
        # 2. 让AI选择院校及专业，并返回对应ID
        ai_selection = ai_select_college_ids(filtered_colleges, user_info, recommendation_data, is_first=is_first)

        # 如果AI没有选择结果，直接返回
        if not ai_selection:
            return plan_id, False

        # 3. 根据AI选择结果构建完整的院校志愿数据
        colleges_data = []
        for idx, (cgid, selected_spids) in enumerate(ai_selection.items()):
            # 将字符串类型的cgid转换为整数
            cgid = int(cgid) if isinstance(cgid, str) and cgid.isdigit() else (cgid if isinstance(cgid, int) else 0)
            if cgid == 0:
                continue
                
            # 查找对应的完整院校数据
            college_data = next((c for c in filtered_colleges if c['cgid'] == cgid), None)
            if not college_data:
                continue
                
            # 计算在整个方案中的序号（1-48）
            volunteer_index = (actual_group_id - 1) * 4 + idx + 1
            
            # 构建院校志愿数据
            college_volunteer = {
                'category_id': category_id,
                'group_id': actual_group_id,
                'volunteer_index': volunteer_index,
                'college_id': college_data['cid'],
                'college_name': college_data['cname'],
                'college_group_id': college_data['cgid'],
                'score_diff': college_data['score_diff'],
                'prediction_score': college_data['min_score'],
                'recommend_type': 'ai',
                'specialties': [],
                'area_name': college_data['area_name'],
                'group_name': college_data['group_name'],
                'min_tuition': college_data['min_tuition'],
                'max_tuition': college_data['max_tuition'],
                'min_score': college_data['min_score'],
                'plan_number': college_data['plan_number'],
                'school_type_text': college_data['school_type_text'],
                'subject_requirements': college_data['subject_requirements'],
                'tese_text': college_data['tese_text'],
                'teshu_text': college_data['teshu_text'],
                'uncode': college_data['uncode'],
                'nature': college_data['school_nature'],  
            }
            
            # 处理专业ID列表
            valid_spids = []
            for spid in selected_spids:
                # 确保spid是有效的整数
                if isinstance(spid, str) and spid.isdigit():
                    valid_spids.append(int(spid))
                elif isinstance(spid, int):
                    valid_spids.append(spid)
            
            # 只添加AI选中的专业数据
            selected_spids_set = set(valid_spids)
            sp_idx = 0
            for specialty in college_data['specialties']:
                if sp_idx >= 6:  # 最多添加6个专业
                    break
                    
                specialty_id = specialty.get('spid', 0)
                if specialty_id in selected_spids_set:
                    # 添加专业数据
                    specialty_data = {
                        'specialty_id': specialty_id,
                        'specialty_code': specialty.get('spcode', ''),
                        'specialty_name': specialty.get('spname', ''),
                        'specialty_index': sp_idx + 1,  # 专业序号从1开始
                        'prediction_score': int(specialty.get('prediction_score', 0) or 0),
                        'plan_number': int(specialty.get('plan_number', 0) or 0),
                        'tuition': int(specialty.get('tuition', 0) or 0),
                        'fenshuxian_id': int(specialty.get('id', 0) or 0)
                    }
                    college_volunteer['specialties'].append(specialty_data)
                    sp_idx += 1
                
            colleges_data.append(college_volunteer)

        # 4. 更新志愿方案
        if not plan_id:
            # 如果没有方案ID，创建新方案
            plan = VolunteerPlanService.create_empty_plan(
                student_id=student_id,
                planner_id=planner_id,
                remarks="AI生成的志愿方案"
            )
            plan_id = plan['id']

        # 批量添加院校志愿和专业
        if colleges_data:
            # 添加院校志愿
            college_result = VolunteerPlanService.batch_add_volunteer_colleges(plan_id, colleges_data)
            
            # 查询刚添加的院校志愿，获取它们的ID
            added_colleges = VolunteerCollege.query.filter_by(plan_id=plan_id).all()
            college_id_map = {college.volunteer_index: college.id for college in added_colleges}
            
            # 为每个院校添加专业志愿
            for college_data in colleges_data:
                volunteer_index = college_data['volunteer_index']
                if volunteer_index in college_id_map:
                    college_id = college_id_map[volunteer_index]
                    specialties_data = college_data.get('specialties', [])
                    if specialties_data:
                        VolunteerPlanService.batch_add_volunteer_specialties(college_id, specialties_data)
            
        return plan_id, bool(colleges_data)  # 返回方案ID和成功状态
        
    except Exception as e:
        current_app.logger.error(f"处理批次志愿失败: {str(e)}")
        # 重新抛出异常或返回失败状态
        return plan_id, False
    
def generate_complete_volunteer_plan(student_id, planner_id, user_data_hash, is_first=False):
    """
    生成完整的志愿方案(包含进度跟踪)
    
    :param student_id: 学生ID
    :param planner_id: 规划师ID
    :return: 生成的志愿方案
    """
    # 获取学生数据快照
    student_snapshot = StudentDataService.generate_student_data_snapshot(student_id)
    current_snapshot = json.dumps(student_snapshot, ensure_ascii=False)

    # 查找之前的方案
    previous_plan = StudentVolunteerPlan.query.filter_by(
        student_id=student_id,
        is_current=True
    ).order_by(StudentVolunteerPlan.version.desc()).first()
    previous_snapshot = json.dumps(previous_plan.student_data_snapshot, ensure_ascii=False) if previous_plan else None

    # 创建空方案并设置所有初始状态
    plan = VolunteerPlanService.create_empty_plan(
        student_id=student_id,
        planner_id=planner_id,
        remarks="AI生成的志愿方案",
        user_data_hash=user_data_hash,
        generation_status=StudentVolunteerPlan.GENERATION_STATUS_PROCESSING,
        generation_progress=0,
        generation_message="开始生成志愿方案",
        student_data_snapshot=current_snapshot
    )
    plan_id = plan['id']

    # 让AI分析两次学生快照的差异，以直观展示历史方案的差异
    from app.tasks.volunteer_tasks import analyze_student_snapshots_ai
    analyze_student_snapshots_ai.delay(plan_id, current_snapshot, previous_snapshot)

    try:
        # 处理所有批次
        batch_count = 12
        processed_count = 0
        
        # 遍历所有类别和组
        for category_id in [1,2,3]:  # 冲、稳、保
            for group_id in range(1, 5):  # 每类4个小组(1-4)
                # 处理一个批次
                plan_id, success = process_batch(
                    student_id=student_id,
                    planner_id=planner_id,
                    category_id=category_id,
                    group_id=group_id,
                    plan_id=plan_id,
                    is_first=is_first
                )
                
                # 更新进度
                processed_count += 1
                progress = int((processed_count / batch_count) * 100)
                
                # 更新方案状态
                StudentVolunteerPlan.query.filter_by(id=plan_id).update({
                    'generation_progress': progress,
                    'generation_message': f"已处理{processed_count}/{batch_count}个批次"
                })
                db.session.commit()
        
        # 全部处理完成，更新状态
        StudentVolunteerPlan.query.filter_by(id=plan_id).update({
            'generation_status': StudentVolunteerPlan.GENERATION_STATUS_SUCCESS,
            'generation_progress': 100,
            'generation_message': "志愿方案生成完成"
        })
        db.session.commit()
        
        # 自动生成总方案和三个类别的分析
        from app.tasks.volunteer_tasks import analyze_volunteer_plan_task, analyze_volunteer_category_task
        
        # 启动总方案分析任务
        analyze_volunteer_plan_task.delay(plan_id)
        
        # 启动三个类别分析任务
        for category_id in [1, 2, 3]:  # 冲、稳、保
            analyze_volunteer_category_task.delay(plan_id, category_id)
        
        # 返回完整方案
        return VolunteerPlanService.get_volunteer_plan(plan_id)
        
    except Exception as e:
        # 发生错误，更新状态
        StudentVolunteerPlan.query.filter_by(id=plan_id).update({
            'generation_status': StudentVolunteerPlan.GENERATION_STATUS_FAILED,
            'generation_message': f"生成失败: {str(e)}"
        })
        db.session.commit()
        raise